import os
import logging
import json
import requests

logger = logging.getLogger("wecom_bot")

def extract_pdf_text(file_path):
    try:
        import pypdf
        reader = pypdf.PdfReader(file_path)
        text = ""
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
        return text.strip()
    except ImportError:
        raise ImportError("pypdf")

def extract_docx_text(file_path):
    try:
        import docx
        doc = docx.Document(file_path)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                text.append(" | ".join(row_text))
        return "\n".join(text).strip()
    except ImportError:
        raise ImportError("python-docx")

def extract_xlsx_text(file_path):
    try:
        import pandas as pd
        # 读取所有 sheet
        xl = pd.ExcelFile(file_path)
        sheets_text = []
        for sheet_name in xl.sheet_names:
            df = xl.parse(sheet_name)
            # 使用 tabulate 渲染成 markdown 风格文本，或者简易 string 渲染
            try:
                sheet_str = df.to_markdown(index=False)
            except ImportError:
                sheet_str = df.to_string(index=False)
            sheets_text.append(f"Sheet: {sheet_name}\n{sheet_str}")
        return "\n\n".join(sheets_text).strip()
    except ImportError:
        raise ImportError("pandas or openpyxl")

def convert_pdf_to_images(file_path, output_dir):
    try:
        from pdf2image import convert_from_path
        pages = convert_from_path(file_path, dpi=150)
        image_paths = []
        for i, page in enumerate(pages):
            image_path = os.path.join(output_dir, f"page_{i}.png")
            page.save(image_path, "PNG")
            image_paths.append(image_path)
        return image_paths
    except Exception as e:
        logger.error(f"Failed to convert PDF to images: {e}")
        return []

def extract_fields_from_text(text: str, doc_type: str) -> dict:
    """使用 DeepSeek 大模型从提取到的文档文本中直接提取结构化 JSON 字段"""
    api_key = os.getenv("DEEPSEEK_API_KEY")
    if not api_key:
        logger.error("DEEPSEEK_API_KEY not configured")
        return {}
        
    prompt = (
        f"你是一个智能审单专家。这是一份类型为 {doc_type} 的单据文本内容。\n"
        f"请仔细分析该文本，提取以下关键字段的值（如果没有找到，请将其值设为 null）：\n"
        f"- type (固定为 {doc_type.upper()})\n"
        f"- net_weight (提取净重数值)\n"
        f"- unit_price (提取单价)\n"
        f"- total_amount (提取总金额)\n"
        f"- latest_shipment_date (提取最晚装运期，格式 YYYY-MM-DD)\n"
        f"- container_no (提取集装箱号，通常格式为4位字母+7位数字)\n"
        f"- seal_no (提取封条号)\n"
        f"- shipper_name (提取发货人/公司名称)\n"
        f"- shipper_address (提取发货人地址)\n"
        f"- consignee_name (提取收货人/公司名称)\n"
        f"- consignee_address (提取收货人地址)\n\n"
        f"单据文本内容：\n"
        f"{text[:6000]}\n\n"
        f"请仅输出 JSON 格式，不要包含任何 markdown 标记（例如 ```json）或其他解释文本。"
    )
    
    try:
        url = "https://api.deepseek.com/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": "deepseek-chat",
            "messages": [
                {"role": "system", "content": "你是一个专业提取文档字段的 AI 助手。只输出纯 JSON，不含格式。"},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.0
        }
        
        resp = requests.post(url, headers=headers, json=payload, timeout=45)
        if resp.status_code == 200:
            content = resp.json()["choices"][0]["message"]["content"]
            clean_json = content.replace("```json", "").replace("```", "").strip()
            return json.loads(clean_json)
        else:
            logger.error(f"DeepSeek text extraction failed: {resp.text}")
            return {}
    except Exception as e:
        logger.error(f"DeepSeek text extraction exception: {e}")
        return {}
